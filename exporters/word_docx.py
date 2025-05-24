from .base import Exporter
from docx import Document
import pandas as pd


class WordExporter(Exporter):
    def export(self, dataframe: pd.DataFrame, output_path: str):
        doc = Document()
        doc.add_heading("Tabela do Memorial SIGEF", level=1)

        registros_texto = []
        for _, row in dataframe.iterrows():
            linha = (
                f"Vértice: {row['Vértice']}, "
                f"Longitude: {row['Longitude']}, "
                f"Latitude: {row['Latitude']}, "
                f"Altitude (m): {row['Altitude (m)']}, "
                f"Vante: {row['Vante']}, "
                f"Azimute: {row['Azimute']}, "
                f"Distância (m): {row['Distância (m)']}, "
                f"Confrontações: {row['Confrontações']}"
            )
            registros_texto.append(linha)

        corpo_final = "; ".join(registros_texto) + ";"
        doc.add_paragraph(corpo_final)
        doc.save(output_path)
        print(f"📄 Word salvo em: {output_path}")
